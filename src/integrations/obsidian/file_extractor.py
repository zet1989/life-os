"""Извлечение текста из файлов различных форматов.

Поддерживаемые форматы:
- PDF (PyMuPDF / fitz)
- DOCX (python-docx)
- DOC (через docx, если реально .docx; иначе best-effort)
- XLSX/XLS (openpyxl)
- CSV (stdlib csv)
- HTML (BeautifulSoup)
- TXT (plain text)
"""

import csv
import io
from pathlib import Path

import structlog

logger = structlog.get_logger()

# Расширения, которые поддерживаются помимо .md
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".html", ".htm", ".txt"}


def extract_text(filepath: str | Path) -> str | None:
    """Извлечь текст из файла. Возвращает None при ошибке или неподдерживаемом формате."""
    path = Path(filepath)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return None

    try:
        extractor = _EXTRACTORS.get(ext)
        if extractor:
            return extractor(path)
    except Exception:
        logger.exception("file_extractor.error", file=str(path), ext=ext)

    return None


def _extract_pdf(path: Path) -> str:
    """Извлечь текст из PDF через PyMuPDF."""
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    pages: list[str] = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text.strip())
    doc.close()
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    """Извлечь текст из DOCX через python-docx."""
    from docx import Document

    doc = Document(str(path))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def _extract_xlsx(path: Path) -> str:
    """Извлечь текст из XLSX через openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    sheets: list[str] = []

    for ws in wb.worksheets:
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cells).strip(" |")
            if line:
                rows.append(line)
        if rows:
            header = f"## {ws.title}" if ws.title else ""
            sheets.append(f"{header}\n" + "\n".join(rows) if header else "\n".join(rows))

    wb.close()
    return "\n\n".join(sheets)


def _extract_csv(path: Path) -> str:
    """Извлечь текст из CSV."""
    # Пробуем определить кодировку
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            raw = path.read_text(encoding=encoding)
            break
        except (UnicodeDecodeError, ValueError):
            continue
    else:
        return ""

    reader = csv.reader(io.StringIO(raw))
    rows: list[str] = []
    for row in reader:
        line = " | ".join(row).strip(" |")
        if line:
            rows.append(line)
    return "\n".join(rows)


def _extract_html(path: Path) -> str:
    """Извлечь текст из HTML через BeautifulSoup."""
    from bs4 import BeautifulSoup

    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            raw = path.read_text(encoding=encoding)
            break
        except (UnicodeDecodeError, ValueError):
            continue
    else:
        return ""

    soup = BeautifulSoup(raw, "html.parser")

    # Убираем скрипты и стили
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    # Убираем лишние пустые строки
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _extract_txt(path: Path) -> str:
    """Извлечь текст из TXT."""
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return ""


# Маппинг расширение → функция-экстрактор
_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".doc": _extract_docx,  # best-effort: если .doc — реально .docx
    ".xlsx": _extract_xlsx,
    ".xls": _extract_xlsx,  # openpyxl может не читать xls; best-effort
    ".csv": _extract_csv,
    ".html": _extract_html,
    ".htm": _extract_html,
    ".txt": _extract_txt,
}
